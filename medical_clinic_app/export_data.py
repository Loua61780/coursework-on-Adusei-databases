import json
import csv
import pandas as pd
from datetime import datetime
from pathlib import Path
from sqlalchemy.orm import joinedload
import os

# Импорт библиотек для PDF и DOCX (установите через pip)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Предупреждение: библиотека reportlab не установлена. Экспорт в PDF будет недоступен.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Предупреждение: библиотека python-docx не установлена. Экспорт в DOCX будет недоступен.")

class DataExporter:
    """Класс для экспорта данных в различные форматы"""
    
    def __init__(self, session):
        self.session = session
        self.export_dir = Path("exports")
        self.export_dir.mkdir(exist_ok=True)
    
    def export_appointments_to_json(self, start_date=None, end_date=None, doctor_id=None):
        """Экспорт записей на прием в JSON"""
        try:
            query = self.session.query(Appointment).options(
                joinedload(Appointment.patient),
                joinedload(Appointment.doctor),
                joinedload(Appointment.schedule)
            )
            
            if start_date:
                query = query.filter(Appointment.appointment_date >= start_date)
            if end_date:
                query = query.filter(Appointment.appointment_date <= end_date)
            if doctor_id:
                query = query.filter(Appointment.doctor_id == doctor_id)
            
            appointments = query.all()
            
            data = []
            for appt in appointments:
                appointment_data = {
                    'id': appt.id,
                    'patient': {
                        'id': appt.patient.id,
                        'full_name': appt.patient.full_name,
                        'phone': appt.patient.phone
                    } if appt.patient else None,
                    'doctor': {
                        'id': appt.doctor.id,
                        'full_name': appt.doctor.full_name,
                        'specialization': appt.doctor.specialization.name if appt.doctor.specialization else None
                    } if appt.doctor else None,
                    'appointment_date': appt.appointment_date.isoformat() if appt.appointment_date else None,
                    'appointment_time': str(appt.appointment_time) if appt.appointment_time else None,
                    'status': appt.status.value if appt.status else None,
                    'reason': appt.reason,
                    'created_at': appt.created_at.isoformat() if appt.created_at else None
                }
                data.append(appointment_data)
            
            # Генерация имени файла с timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.export_dir / f'appointments_{timestamp}.json'
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2, default=str)
            
            return True, f"Данные экспортированы в {filename}", str(filename)
        except Exception as e:
            return False, f"Ошибка экспорта в JSON: {str(e)}", None
    
    def export_patients_to_csv(self, min_age=None, max_age=None):
        """Экспорт пациентов в CSV"""
        try:
            from models import Patient
            query = self.session.query(Patient)
            
            patients = query.all()
            
            # Подготовка данных
            data = []
            for patient in patients:
                patient_data = {
                    'id': patient.id,
                    'last_name': patient.last_name,
                    'first_name': patient.first_name,
                    'patronymic': patient.patronymic,
                    'birth_date': patient.birth_date.isoformat() if patient.birth_date else None,
                    'age': patient.age,
                    'gender': patient.gender,
                    'phone': patient.phone,
                    'email': patient.email,
                    'address': patient.address,
                    'registration_date': patient.registration_date.isoformat() if patient.registration_date else None
                }
                data.append(patient_data)
            
            # Фильтрация по возрасту если задано
            if min_age is not None or max_age is not None:
                filtered_data = []
                for patient in data:
                    if patient['age'] is not None:
                        if min_age is not None and patient['age'] < min_age:
                            continue
                        if max_age is not None and patient['age'] > max_age:
                            continue
                        filtered_data.append(patient)
                data = filtered_data
            
            # Создание DataFrame и экспорт в CSV
            df = pd.DataFrame(data)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.export_dir / f'patients_{timestamp}.csv'
            
            df.to_csv(filename, index=False, encoding='utf-8-sig')
            
            return True, f"Данные экспортированы в {filename}", str(filename)
        except Exception as e:
            return False, f"Ошибка экспорта в CSV: {str(e)}", None
    
    def export_schedule_to_pdf(self, doctor_id=None, start_date=None, end_date=None):
        """Экспорт расписания в PDF"""
        if not PDF_AVAILABLE:
            return False, "Библиотека reportlab не установлена", None
        
        try:
            from models import Schedule, Employee
            
            query = self.session.query(Schedule).options(
                joinedload(Schedule.employee).joinedload(Employee.specialization)
            )
            
            if doctor_id:
                query = query.filter(Schedule.employee_id == doctor_id)
            if start_date:
                query = query.filter(Schedule.work_date >= start_date)
            if end_date:
                query = query.filter(Schedule.work_date <= end_date)
            
            schedules = query.order_by(Schedule.work_date, Schedule.start_time).all()
            
            if not schedules:
                return False, "Нет данных для экспорта", None
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.export_dir / f'schedule_{timestamp}.pdf'
            
            # Создание PDF документа
            doc = SimpleDocTemplate(str(filename), pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Заголовок
            title = Paragraph("Расписание врачей", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 12))
            
            # Подготовка данных для таблицы
            table_data = [['Дата', 'Врач', 'Специализация', 'Время приема', 'Кабинет', 'Свободных мест']]
            
            for schedule in schedules:
                doctor_name = schedule.employee.full_name if schedule.employee else "Не указан"
                specialization = schedule.employee.specialization.name if schedule.employee and schedule.employee.specialization else "Не указана"
                
                table_data.append([
                    schedule.work_date.strftime("%d.%m.%Y") if schedule.work_date else "",
                    doctor_name,
                    specialization,
                    f"{schedule.start_time.strftime('%H:%M')} - {schedule.end_time.strftime('%H:%M')}",
                    schedule.cabinet_number or "",
                    str(schedule.available_slots)
                ])
            
            # Создание таблицы
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            elements.append(table)
            
            # Добавление информации о дате экспорта
            elements.append(Spacer(1, 20))
            export_date = Paragraph(f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", styles['Normal'])
            elements.append(export_date)
            
            # Сборка документа
            doc.build(elements)
            
            return True, f"Данные экспортированы в {filename}", str(filename)
        except Exception as e:
            return False, f"Ошибка экспорта в PDF: {str(e)}", None
    
    def export_medical_records_to_docx(self, patient_id=None, doctor_id=None):
        """Экспорт медицинских записей в DOCX"""
        if not DOCX_AVAILABLE:
            return False, "Библиотека python-docx не установлена", None
        
        try:
            from models import MedicalRecord, Patient, Employee, Diagnosis
            
            query = self.session.query(MedicalRecord).options(
                joinedload(MedicalRecord.patient),
                joinedload(MedicalRecord.doctor),
                joinedload(MedicalRecord.diagnosis)
            )
            
            if patient_id:
                query = query.filter(MedicalRecord.patient_id == patient_id)
            if doctor_id:
                query = query.filter(MedicalRecord.doctor_id == doctor_id)
            
            records = query.order_by(MedicalRecord.record_date.desc()).all()
            
            if not records:
                return False, "Нет данных для экспорта", None
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.export_dir / f'medical_records_{timestamp}.docx'
            
            # Создание DOCX документа
            doc = Document()
            
            # Заголовок
            doc.add_heading('Медицинские записи', 0)
            
            # Добавление информации о фильтрах
            if patient_id:
                patient = self.session.query(Patient).filter_by(id=patient_id).first()
                if patient:
                    doc.add_paragraph(f'Пациент: {patient.full_name}')
            
            if doctor_id:
                doctor = self.session.query(Employee).filter_by(id=doctor_id).first()
                if doctor:
                    doc.add_paragraph(f'Врач: {doctor.full_name}')
            
            doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
            doc.add_paragraph()
            
            # Добавление записей
            for i, record in enumerate(records, 1):
                doc.add_heading(f'Запись #{i}', level=2)
                
                # Информация о приеме
                doc.add_paragraph(f'Дата приема: {record.record_date.strftime("%d.%m.%Y %H:%M") if record.record_date else "Не указана"}')
                doc.add_paragraph(f'Пациент: {record.patient.full_name if record.patient else "Не указан"}')
                doc.add_paragraph(f'Врач: {record.doctor.full_name if record.doctor else "Не указан"}')
                
                if record.diagnosis:
                    doc.add_paragraph(f'Диагноз: {record.diagnosis.code} - {record.diagnosis.name}')
                
                # Жалобы
                if record.complaints:
                    doc.add_heading('Жалобы', level=3)
                    doc.add_paragraph(record.complaints)
                
                # Результаты осмотра
                if record.examination_results:
                    doc.add_heading('Результаты осмотра', level=3)
                    doc.add_paragraph(record.examination_results)
                
                # Рекомендации
                if record.recommendations:
                    doc.add_heading('Рекомендации', level=3)
                    doc.add_paragraph(record.recommendations)
                
                doc.add_paragraph('-' * 50)
            
            # Сохранение документа
            doc.save(str(filename))
            
            return True, f"Данные экспортированы в {filename}", str(filename)
        except Exception as e:
            return False, f"Ошибка экспорта в DOCX: {str(e)}", None
    
    def export_statistics_to_xlsx(self):
        """Экспорт статистики в XLSX"""
        try:
            from models import Patient, Employee, Appointment, MedicalRecord
            
            # Сбор данных для разных листов
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.export_dir / f'statistics_{timestamp}.xlsx'
            
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                # Лист с пациентами
                patients = self.session.query(Patient).all()
                patients_data = []
                for p in patients:
                    patients_data.append({
                        'ID': p.id,
                        'Фамилия': p.last_name,
                        'Имя': p.first_name,
                        'Отчество': p.patronymic or '',
                        'Дата рождения': p.birth_date,
                        'Возраст': p.age,
                        'Телефон': p.phone or '',
                        'Email': p.email or '',
                        'Дата регистрации': p.registration_date
                    })
                
                if patients_data:
                    df_patients = pd.DataFrame(patients_data)
                    df_patients.to_excel(writer, sheet_name='Пациенты', index=False)
                
                # Лист с врачами
                doctors = self.session.query(Employee).join(Position).filter(Position.name.ilike('%врач%')).all()
                doctors_data = []
                for d in doctors:
                    doctors_data.append({
                        'ID': d.id,
                        'Фамилия': d.last_name,
                        'Имя': d.first_name,
                        'Отчество': d.patronymic or '',
                        'Должность': d.position.name if d.position else '',
                        'Специализация': d.specialization.name if d.specialization else '',
                        'Кабинет': d.cabinet_number or '',
                        'Телефон': d.phone or '',
                        'Email': d.email or ''
                    })
                
                if doctors_data:
                    df_doctors = pd.DataFrame(doctors_data)
                    df_doctors.to_excel(writer, sheet_name='Врачи', index=False)
                
                # Лист со статистикой
                stats_data = []
                
                # Количество пациентов
                total_patients = self.session.query(Patient).count()
                stats_data.append(['Общее количество пациентов', total_patients])
                
                # Количество врачей
                total_doctors = self.session.query(Employee).join(Position).filter(Position.name.ilike('%врач%')).count()
                stats_data.append(['Количество врачей', total_doctors])
                
                # Количество записей
                total_appointments = self.session.query(Appointment).count()
                stats_data.append(['Общее количество записей', total_appointments])
                
                # Записи по статусам
                status_counts = self.session.query(Appointment.status, func.count(Appointment.id)).group_by(Appointment.status).all()
                for status, count in status_counts:
                    stats_data.append([f'Записи со статусом "{status.value}"', count])
                
                # Количество медицинских записей
                total_records = self.session.query(MedicalRecord).count()
                stats_data.append(['Количество медицинских записей', total_records])
                
                df_stats = pd.DataFrame(stats_data, columns=['Показатель', 'Значение'])
                df_stats.to_excel(writer, sheet_name='Статистика', index=False)
                
                # Лист с последними записями
                recent_appointments = self.session.query(Appointment).order_by(Appointment.created_at.desc()).limit(50).all()
                appointments_data = []
                for a in recent_appointments:
                    appointments_data.append({
                        'ID': a.id,
                        'Дата приема': a.appointment_date,
                        'Время приема': a.appointment_time,
                        'Пациент': a.patient.full_name if a.patient else '',
                        'Врач': a.doctor.full_name if a.doctor else '',
                        'Статус': a.status.value if a.status else '',
                        'Причина': a.reason or ''
                    })
                
                if appointments_data:
                    df_appointments = pd.DataFrame(appointments_data)
                    df_appointments.to_excel(writer, sheet_name='Последние записи', index=False)
            
            return True, f"Данные экспортированы в {filename}", str(filename)
        except Exception as e:
            return False, f"Ошибка экспорта в XLSX: {str(e)}", None
    
    def export_all_formats(self):
        """Экспорт данных во все доступные форматы"""
        results = []
        
        # Экспорт в JSON
        success, message, filepath = self.export_appointments_to_json()
        results.append(('JSON', success, message, filepath))
        
        # Экспорт в CSV
        success, message, filepath = self.export_patients_to_csv()
        results.append(('CSV', success, message, filepath))
        
        # Экспорт в PDF
        if PDF_AVAILABLE:
            success, message, filepath = self.export_schedule_to_pdf()
            results.append(('PDF', success, message, filepath))
        else:
            results.append(('PDF', False, 'Библиотека reportlab не установлена', None))
        
        # Экспорт в DOCX
        if DOCX_AVAILABLE:
            success, message, filepath = self.export_medical_records_to_docx()
            results.append(('DOCX', success, message, filepath))
        else:
            results.append(('DOCX', False, 'Библиотека python-docx не установлена', None))
        
        # Экспорт в XLSX
        success, message, filepath = self.export_statistics_to_xlsx()
        results.append(('XLSX', success, message, filepath))
        
        return results